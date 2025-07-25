import io
from PyPDF2 import PdfReader
import docx

class CVParser:
    @staticmethod
    async def parse(file):
        content = await file.read()
        if file.filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        elif file.filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])
        elif file.filename.endswith(".txt"):
            return content.decode("utf-8")
        else:
            return "지원하지 않는 파일 형식입니다."

async def extract_text_from_file(file):
    """파일에서 텍스트 추출하는 함수"""
    content = await file.read()
    
    if file.filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([p.text for p in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        return content.decode("utf-8")
    else:
        raise ValueError("지원하지 않는 파일 형식입니다. PDF, DOCX, TXT 파일만 지원합니다.")